"""
报告导出模块 - 生成 Word 文档
将合同风险分析报告导出为专业格式的 Word 文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
from typing import List, Dict, Any
from datetime import datetime
import os


def set_chinese_font(run, font_name="微软雅黑", font_size=12):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def generate_word_report(
    structured_data: List[Dict[str, Any]],
    report_md: str,
    risk_score: int,
    output_dir: Path = None,
    filename: str = None
) -> str:
    """
    生成 Word 格式的风险分析报告
    
    Args:
        structured_data: 结构化风险数据列表
        report_md: Markdown 格式的摘要报告
        risk_score: 风险评分
        output_dir: 输出目录
        filename: 输出文件名
    
    Returns:
        生成的 Word 文件路径
    """
    if output_dir is None:
        output_dir = Path(__file__).parent.parent.parent / "temp_reports"
    
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    if filename is None:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"合同风险分析报告_{timestamp}.docx"
    
    # 创建文档
    doc = Document()
    
    # ========== 封面标题 ==========
    title = doc.add_heading("合同风险分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run, "微软雅黑", 28)
        run.font.bold = True
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}")
    set_chinese_font(run, "微软雅黑", 12)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()  # 空行
    
    # ========== 风险评分概览 ==========
    doc.add_heading("一、风险评分概览", level=1)
    
    # 风险等级判定
    if risk_score < 30:
        risk_level = "低风险"
        risk_color = RGBColor(16, 185, 129)  # 绿色
        risk_advice = "合同整体风险可控，建议关注标注条款后可正常签署。"
    elif risk_score < 70:
        risk_level = "中风险"
        risk_color = RGBColor(245, 158, 11)  # 黄色
        risk_advice = "合同存在一定风险，建议对高风险条款进行修改后再签署。"
    else:
        risk_level = "高风险"
        risk_color = RGBColor(239, 68, 68)  # 红色
        risk_advice = "合同存在重大风险隐患，强烈建议修改相关条款或寻求专业法律意见。"
    
    # 统计数据
    high_count = sum(1 for item in structured_data if item.get("risk_level") == "高")
    medium_count = sum(1 for item in structured_data if item.get("risk_level") == "中")
    low_count = sum(1 for item in structured_data if item.get("risk_level") == "低")
    total_count = len(structured_data)
    
    # 评分表格
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    rows_data = [
        ("综合风险评分", f"{risk_score} 分"),
        ("风险等级", risk_level),
        ("检测条款", f"{total_count} 处"),
        ("风险分布", f"高风险 {high_count} 处 / 中风险 {medium_count} 处 / 低风险 {low_count} 处"),
    ]
    
    for i, (label, value) in enumerate(rows_data):
        cells = table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, "微软雅黑", 11)
    
    doc.add_paragraph()
    
    # 风险建议
    advice_para = doc.add_paragraph()
    advice_run = advice_para.add_run(f"📋 审查建议：{risk_advice}")
    set_chinese_font(advice_run, "微软雅黑", 11)
    advice_run.font.bold = True
    
    doc.add_paragraph()
    
    # ========== 条款详细分析 ==========
    doc.add_heading("二、条款详细分析", level=1)
    
    for idx, item in enumerate(structured_data):
        risk_level = item.get("risk_level", "低")
        is_high_risk = risk_level == "高"
        
        # 条款标题
        clause_heading = doc.add_heading(f"条款 {idx + 1}", level=2)
        for run in clause_heading.runs:
            set_chinese_font(run, "微软雅黑", 14)
        
        # 风险等级标签（高/中/低三级）
        risk_para = doc.add_paragraph()
        if risk_level == "高":
            risk_run = risk_para.add_run("【高风险】")
            risk_run.font.color.rgb = RGBColor(239, 68, 68)  # 红色
        elif risk_level == "中":
            risk_run = risk_para.add_run("【中风险】")
            risk_run.font.color.rgb = RGBColor(249, 115, 22)  # 橙色
        else:
            risk_run = risk_para.add_run("【低风险】")
            risk_run.font.color.rgb = RGBColor(16, 185, 129)  # 绿色
        set_chinese_font(risk_run, "微软雅黑", 11)
        risk_run.font.bold = True
        
        # 原文
        doc.add_paragraph("■ 条款原文：", style='Intense Quote')
        clause_text = doc.add_paragraph(item.get("clause_text", ""))
        clause_text.paragraph_format.left_indent = Inches(0.5)
        for run in clause_text.runs:
            set_chinese_font(run, "宋体", 10)
        
        # 风险分析（仅高风险）
        if is_high_risk:
            doc.add_paragraph("■ 风险分析：", style='Intense Quote')
            analysis = item.get("deep_analysis") or item.get("risk_reason", "")
            analysis_para = doc.add_paragraph(analysis)
            analysis_para.paragraph_format.left_indent = Inches(0.5)
            for run in analysis_para.runs:
                set_chinese_font(run, "微软雅黑", 10)
        
        # 修改建议
        doc.add_paragraph("■ 修改建议：", style='Intense Quote')
        suggestion = item.get("suggestion", "无须修改")
        suggestion_para = doc.add_paragraph(suggestion)
        suggestion_para.paragraph_format.left_indent = Inches(0.5)
        for run in suggestion_para.runs:
            set_chinese_font(run, "微软雅黑", 10)
            run.font.color.rgb = RGBColor(5, 150, 105)  # 绿色
        
        # 法律依据（仅高风险）
        if is_high_risk:
            law_content = item.get("law_content") or item.get("law_reference", "")
            if law_content:
                doc.add_paragraph("■ 法律依据：", style='Intense Quote')
                law_para = doc.add_paragraph(law_content)
                law_para.paragraph_format.left_indent = Inches(0.5)
                for run in law_para.runs:
                    set_chinese_font(run, "宋体", 9)
                    run.font.color.rgb = RGBColor(99, 102, 241)
        
        doc.add_paragraph()  # 条款间空行
    
    # ========== 免责声明 ==========
    doc.add_heading("三、免责声明", level=1)
    
    disclaimer_text = (
        "本报告由 AI 智能系统自动生成，仅供参考，不构成任何法律意见。"
        "在签署重要合同前，请务必咨询专业律师。"
        "报告生成方对因使用本报告内容而产生的任何后果不承担法律责任。"
    )
    disclaimer_para = doc.add_paragraph(disclaimer_text)
    for run in disclaimer_para.runs:
        set_chinese_font(run, "微软雅黑", 9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # ========== 保存文档 ==========
    output_path = output_dir / filename
    doc.save(str(output_path))
    
    print(f"📄 Word 报告已生成: {output_path}")
    return str(output_path)


if __name__ == "__main__":
    # 测试
    test_data = [
        {
            "clause_text": "甲方有权在任何情况下单方解除本合同，无需承担任何责任。",
            "risk_level": "高",
            "risk_reason": "单方解除权条款",
            "deep_analysis": "该条款赋予甲方无条件单方解除权，违反了合同公平原则。",
            "suggestion": "建议修改为: '经双方协商一致, 任一方可解除本合同'。",
            "law_content": "《民法典》第563条：当事人可以约定解除合同的条件。",
        },
        {
            "clause_text": "本合同自双方签字盖章之日起生效。",
            "risk_level": "低",
            "risk_reason": "标准生效条款",
            "suggestion": "无须修改",
        },
    ]
    
    path = generate_word_report(test_data, "", 45)
    print(f"测试完成: {path}")
