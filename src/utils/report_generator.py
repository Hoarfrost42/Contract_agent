import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

def generate_word_report(markdown_text: str) -> BytesIO:
    """
    Convert Markdown text to a formatted Word document (.docx).
    """
    doc = Document()
    
    # --- Style Configuration ---
    # Set default font to Microsoft YaHei (微软雅黑)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Patch for East Asian fonts
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Helper to set heading style
    def set_heading_style(heading, level):
        run = heading.runs[0]
        run.font.name = 'Microsoft YaHei'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.color.rgb = RGBColor(46, 84, 165) # Deep Blue
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(15)
            run.font.bold = True
        elif level == 3:
            run.font.size = Pt(13)
            run.font.bold = True

    # Helper to process bold text (**text**)
    def add_formatted_paragraph(paragraph, text):
        # Split by bold markers
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                content = part[2:-2]
                run = paragraph.add_run(content)
                run.font.bold = True
                run.font.name = 'Microsoft YaHei'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                if part:
                    run = paragraph.add_run(part)
                    run.font.name = 'Microsoft YaHei'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # --- Parsing Markdown ---
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            set_heading_style(h, 1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            set_heading_style(h, 2)
        elif line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            set_heading_style(h, 3)
            
        # List Items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_paragraph(p, line[2:])
            
        # Quotes
        elif line.startswith('> '):
            p = doc.add_paragraph(style='Quote')
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100) # Gray
            
        # Normal Text
        else:
            p = doc.add_paragraph()
            add_formatted_paragraph(p, line)

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
